"""
DOCX Export Tool for Research Reports.

Converts markdown reports to DOCX format and saves them to the output directory.
"""

import re
import logging
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class DOCXExporter:
    """
    Tool for exporting markdown reports to DOCX format.
    
    Uses python-docx to convert markdown to DOCX.
    """

    def __init__(self, output_dir: str = "output") -> None:
        """
        Initialize the DOCX exporter.
        
        Args:
            output_dir: Directory where DOCX files will be saved
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        if Document is None:
            logger.warning(
                "DOCXExporter: python-docx not installed. "
                "Install with: pip install python-docx"
            )

    def sanitize_filename(self, filename: str) -> str:
        """
        Sanitize a string to be used as a filename.
        
        Removes or replaces invalid characters for filenames.
        """
        # Remove or replace invalid characters
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        # Remove leading/trailing spaces and dots
        filename = filename.strip(' .')
        # Limit length
        if len(filename) > 200:
            filename = filename[:200]
        # If empty after sanitization, use default
        if not filename:
            filename = "research_report"
        return filename

    def _add_formatted_text(self, paragraph, text: str):
        """
        Add text to a paragraph, handling markdown formatting (bold, italic, links).
        
        Args:
            paragraph: docx paragraph object
            text: Text with markdown formatting
        """
        # Pattern to match markdown formatting: **bold**, *italic*, [link](url)
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))', text)
        
        for part in parts:
            if not part:
                continue
            
            # Bold text: **text**
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            # Italic text: *text*
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            # Link: [text](url)
            elif part.startswith('[') and '](' in part:
                match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if match:
                    link_text = match.group(1)
                    link_url = match.group(2)
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
                    run.underline = True
                    # Add hyperlink (requires additional processing)
                    # For now, just add the text with URL in parentheses
                    paragraph.add_run(f' ({link_url})')
            else:
                paragraph.add_run(part)

    def _parse_markdown_to_docx(self, doc: Document, markdown_content: str):
        """
        Parse markdown content and add it to the DOCX document.
        
        Args:
            doc: docx Document object
            markdown_content: Markdown formatted text
        """
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Empty line
            if not line:
                doc.add_paragraph()
                i += 1
                continue
            
            # Header level 1: # Header
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.style.font.size = Pt(18)
                i += 1
                continue
            
            # Header level 2: ## Header
            if line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
                heading.style.font.size = Pt(16)
                i += 1
                continue
            
            # Header level 3: ### Header
            if line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
                heading.style.font.size = Pt(14)
                i += 1
                continue
            
            # Unordered list: - item or * item
            if line.startswith('- ') or line.startswith('* '):
                item_text = line[2:].strip()
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, item_text)
                i += 1
                continue
            
            # Ordered list: 1. item
            if re.match(r'^\d+\.\s+', line):
                item_text = re.sub(r'^\d+\.\s+', '', line)
                para = doc.add_paragraph(style='List Number')
                self._add_formatted_text(para, item_text)
                i += 1
                continue
            
            # Code block: ```code```
            if line.startswith('```'):
                # Collect code block
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    i += 1  # Skip closing ```
                
                # Add code block as a paragraph with monospace font
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                para.style = 'No Spacing'
                continue
            
            # Regular paragraph
            para = doc.add_paragraph()
            self._add_formatted_text(para, line)
            i += 1

    def export_to_docx(
        self,
        markdown_content: str,
        user_query: str,
        filename: Optional[str] = None
    ) -> str:
        """
        Export a markdown report to DOCX.
        
        Args:
            markdown_content: Markdown formatted report
            user_query: Original user query (used for filename if filename not provided)
            filename: Optional custom filename (without .docx extension)
            
        Returns:
            Path to the saved DOCX file
            
        Raises:
            ImportError: If required libraries are not installed
            Exception: If DOCX generation fails
        """
        if Document is None:
            raise ImportError(
                "python-docx library is not installed. Install with: pip install python-docx"
            )
        
        # Determine filename
        if filename is None:
            filename = self.sanitize_filename(user_query)
        else:
            filename = self.sanitize_filename(filename)
        
        # Ensure .docx extension
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        docx_path = self.output_dir / filename
        
        try:
            logger.info("DOCXExporter: Creating DOCX document...")
            
            # Create new document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Parse and add markdown content
            logger.info("DOCXExporter: Parsing markdown content...")
            self._parse_markdown_to_docx(doc, markdown_content)
            
            # Save document
            logger.info("DOCXExporter: Saving DOCX to: %s", docx_path)
            doc.save(str(docx_path))
            
            logger.info("DOCXExporter: Successfully saved DOCX to: %s", docx_path)
            return str(docx_path)
            
        except Exception as exc:
            logger.error("DOCXExporter: Failed to generate DOCX: %s", exc)
            raise


def export_report_to_docx(
    markdown_content: str,
    user_query: str,
    output_dir: str = "output",
    filename: Optional[str] = None
) -> str:
    """
    Convenience function to export a markdown report to DOCX.
    
    Args:
        markdown_content: Markdown formatted report
        user_query: Original user query
        output_dir: Directory where DOCX will be saved
        filename: Optional custom filename (without .docx extension)
        
    Returns:
        Path to the saved DOCX file
    """
    exporter = DOCXExporter(output_dir=output_dir)
    return exporter.export_to_docx(markdown_content, user_query, filename)
